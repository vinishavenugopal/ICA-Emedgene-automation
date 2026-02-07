import os
import argparse
import pandas as pd
from docx import Document

# Argument parser to get the root directory from user input
parser = argparse.ArgumentParser(description='Process coverage QC Excel files.')
parser.add_argument(
    '-d', '--directory',
    required=True,
    help='Root directory containing the Excel files'
)
args = parser.parse_args()
root_dir = args.directory

panel_info = {}    # Holds (panel_name, %>=20x)
gene_lists = {}    # Holds [(gene_name, italic), ...]

for subdir, _, files in os.walk(root_dir):
    for file in files:
        if file.endswith('.xlsx') and not file.startswith('~$'):
            file_path = os.path.join(subdir, file)

            # Read from Panel Coverage sheet: %>=20x (E1) and Panel Name (B1)
            try:
                df_panel = pd.read_excel(
                    file_path,
                    sheet_name='Panel Coverage',
                    usecols='B,E',
                    nrows=1,
                    engine='openpyxl'
                )
                panel_name = df_panel.iloc[0, 0]
                value = df_panel.iloc[0, 1]

                if isinstance(value, str) and '%' in value:
                    value = value.replace('%', '').strip()
                value = float(value) * 100  # Convert to percentage

                panel_info[file] = (panel_name, value)

            except Exception as e:
                print(f"Error reading Panel Coverage in {file_path}: {e}")
                panel_info[file] = ("N/A", "N/A")
                continue

            # Read Gene Coverage sheet
            try:
                df_gene = pd.read_excel(
                    file_path,
                    sheet_name='Gene Coverage',
                    engine='openpyxl'
                )

                genes = df_gene.iloc[0:, 0].tolist()
                coverages = df_gene.iloc[0:, 4].tolist()

                genes_with_format = []
                for gene, cov in zip(genes, coverages):
                    if pd.isna(gene):
                        continue
                    try:
                        cov_val = float(cov)
                    except (ValueError, TypeError):
                        cov_val = 100

                    gene_str = str(gene)
                    if cov_val < 0.95:
                        genes_with_format.append((gene_str + '*', True))  # Low coverage
                    else:
                        genes_with_format.append((gene_str, True))  # Italic regardless

                # Sort genes alphabetically
                gene_lists[file] = sorted(genes_with_format, key=lambda x: x[0].lower())

            except Exception as e:
                print(f"Error reading Gene Coverage in {file_path}: {e}")
                gene_lists[file] = []

# Create Word document with table and sorted gene list
doc = Document()
doc.add_heading('Gene Coverage Summary', 0)

for file in sorted(gene_lists.keys()):
    panel_name, coverage_pct = panel_info.get(file, ("N/A", "N/A"))

    doc.add_heading(f"File: {file}", level=1)

    # Add summary table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light List'
    table.cell(0, 0).text = 'Panel Coverage'
    table.cell(0, 1).text = '% of bases ≥ 20x'
    table.cell(1, 0).text = str(panel_name)
    table.cell(1, 1).text = f"{coverage_pct:.2f}%" if isinstance(coverage_pct, (int, float)) else str(coverage_pct)

    doc.add_paragraph()

    # Add sorted gene list (all italic)
    p = doc.add_paragraph()
    for idx, (gene_name, italic) in enumerate(gene_lists[file]):
        run = p.add_run(gene_name)
        run.italic = True  # Always italic
        if idx < len(gene_lists[file]) - 1:
            p.add_run(', ')
    doc.add_paragraph()

# Add explanatory note at the end
doc.add_paragraph(
    "* This gene has suboptimal coverage, defined as less than 95% of its target nucleotides "
    "covered at >20x with a mapping quality score of twenty (MQ≥20). This may be due to "
    "inherent sequencing chemistry limitations or regions of the gene containing duplicated "
    "sequences within the genome."
)

# Save Word output
word_out = os.path.join(root_dir, 'gene_coverage_summary.docx')
doc.save(word_out)
print(f"Saved gene coverage Word document: {word_out}")

